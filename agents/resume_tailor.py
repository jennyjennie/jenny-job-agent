import json
import re
import shutil
from datetime import datetime, timezone
from pathlib import Path
import anthropic
import pdfplumber

from config.settings import Settings

_SYSTEM_PROMPT = """You are a resume expert helping Jenny Ho tailor her resume for a specific job.

Jenny's background:
- UCLA MEng AI (graduating Aug 2026), NYCU BS CS
- Key papers: NeurIPS 2024 first author "Stutter Makes Smarter" (small LLM self-improvement via backward planning);
  ACL 2025 co-author (multimodal text recognition with LLMs)
- MediaTek Research (May 2024 – May 2025): LLM training pipeline optimization, 20% GPU cost reduction,
  evaluation dashboards, "Stutter Makes Smarter" research, Generative Fusion Decoding
- Google Pixel Recorder Team (June 2023 – Sept 2023): QA tooling, Python automation, UI testing
- Gamania Digital Entertainment (July 2022 – Aug 2022): recommendation pipeline
- Project: full-stack AI financial assistant using Claude API tool-calling
- Skills: Python, PyTorch, LLM fine-tuning, inference optimization, FastAPI, Docker, HuggingFace

Your task: produce tailored content for Jenny's resume.

IMPORTANT constraints (to keep resume to exactly ONE page):
- summary: max 2 sentences, max 180 chars total
- MediaTek bullets: exactly 5 bullets, max 130 chars each
- Google bullets: exactly 3 bullets, max 130 chars each
- Gamania bullets: exactly 1 bullet, max 130 chars each

Guidelines:
- Lead with NeurIPS/ACL papers when role is ML/research oriented
- Lead with MediaTek LLM pipeline when role is ML engineering oriented
- Lead with Google when role is SWE/infra oriented
- Use strong action verbs, include metrics wherever possible
- Reorder bullets so the most JD-relevant ones appear first

Return ONLY valid JSON (no prose, no markdown fences):
{
  "missing_keywords": ["keyword1", "keyword2"],
  "tailored_bullets": ["• Action verb ... [metric] ...", "..."],
  "papers_to_highlight": ["NeurIPS 2024 - Stutter Makes Smarter"],
  "summary_suggestion": "2-sentence tailored summary (max 180 chars)",
  "resume_sections": {
    "summary": "2-sentence tailored professional summary for this specific role (max 180 chars)",
    "experience": [
      {
        "company": "MediaTek Research",
        "bullets": [
          "Bullet 1 (max 130 chars)",
          "Bullet 2 (max 130 chars)",
          "Bullet 3 (max 130 chars)",
          "Bullet 4 (max 130 chars)",
          "Bullet 5 (max 130 chars)"
        ]
      },
      {
        "company": "Google",
        "bullets": [
          "Bullet 1 (max 130 chars)",
          "Bullet 2 (max 130 chars)",
          "Bullet 3 (max 130 chars)"
        ]
      },
      {
        "company": "Gamania Digital Entertainment",
        "bullets": [
          "Bullet 1 (max 130 chars)"
        ]
      }
    ]
  }
}"""


def extract_resume_text(pdf_path: str) -> str:
    with pdfplumber.open(pdf_path) as pdf:
        pages = [page.extract_text() or "" for page in pdf.pages]
    return "\n".join(pages)


def tailor_for_job(resume_text: str, job: dict, settings: Settings) -> dict:
    title = job.get("title", "")
    company = job.get("company", "")
    description = (job.get("description", "") or "")[:5000]

    client = anthropic.Anthropic(api_key=settings.anthropic_api_key)
    user_prompt = f"""Target Role: {title} at {company}

Job Description:
{description}

Jenny's Current Resume:
{resume_text[:4000]}"""

    response = client.messages.create(
        model=settings.claude_model,
        max_tokens=3000,
        system=_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_prompt}],
    )
    raw = response.content[0].text.strip()

    # Strip markdown fences if present
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", raw, re.DOTALL)
        if match:
            return json.loads(match.group())
        raise


def _slugify(text: str, max_len: int = 25) -> str:
    """'ML Engineer' → 'MLEngineer', 'Google DeepMind' → 'GoogleDeepMind'"""
    return re.sub(r"[^\w]", "", text.replace(" ", ""))[:max_len]


def _replace_paragraph_text(p, new_text: str):
    """Replace paragraph text in-place, preserving paragraph and first-run formatting."""
    if not p.runs:
        p.add_run(new_text)
        return
    p.runs[0].text = new_text
    for run in p.runs[1:]:
        run.text = ""


def _is_section_header(p) -> bool:
    text = p.text.strip()
    return text in {
        "EDUCATION", "SKILLS", "PUBLICATIONS", "WORK EXPERIENCE", "PROJECTS", "SUMMARY"
    }


def _is_company_header(p) -> bool:
    """Company header: first non-empty run is bold."""
    runs = [r for r in p.runs if r.text.strip()]
    return bool(runs and runs[0].bold)


def _is_title_line(p) -> bool:
    """Title line: italic, not bold (e.g. 'Software Engineer Intern, Pixel Recorder Team')."""
    runs = [r for r in p.runs if r.text.strip()]
    return bool(runs and all(r.italic and not r.bold for r in runs))


def build_docx(result: dict, job: dict, output_dir: str, template_path: str = "") -> str:
    """Build tailored .docx by modifying the base resume template in-place.

    Returns the output file path, or '' if template not found or no resume_sections.
    """
    sections = result.get("resume_sections")
    if not sections:
        return ""

    if not template_path or not Path(template_path).exists():
        print(f"[Tailor] .docx template not found at {template_path!r} — skipping docx")
        return ""

    from docx import Document

    date_str = datetime.now(timezone.utc).strftime("%Y%m%d")
    company_slug = _slugify(job.get("company", "Company"))
    title_slug = _slugify(job.get("title", "Role"))
    filename = f"{date_str}_{company_slug}_{title_slug}.docx"
    output_path = Path(output_dir) / "tailored_resumes" / filename
    output_path.parent.mkdir(parents=True, exist_ok=True)

    shutil.copy2(template_path, str(output_path))
    doc = Document(str(output_path))
    paras = doc.paragraphs

    # ── 1. Replace summary ──────────────────────────────────────────────────
    summary = sections.get("summary", "")
    if summary:
        for i, p in enumerate(paras):
            if i < 2:
                continue
            if _is_section_header(p):
                break
            if p.text.strip() and not _is_company_header(p) and not _is_title_line(p):
                _replace_paragraph_text(p, summary)
                break

    # ── 2. Build company → bullets map ─────────────────────────────────────
    exp_map: dict[str, list[str]] = {}
    for exp in sections.get("experience", []):
        co = exp.get("company", "")
        if co:
            exp_map[co.lower()] = list(exp.get("bullets", []))

    # ── 3. Walk WORK EXPERIENCE section and replace bullets ─────────────────
    in_work_exp = False
    current_bullets: list[str] | None = None
    skip_title = False

    for p in paras:
        text = p.text.strip()

        if not text:
            continue

        if "WORK EXPERIENCE" in text and _is_section_header(p):
            in_work_exp = True
            continue

        if in_work_exp and _is_section_header(p):
            break  # left work experience section

        if not in_work_exp:
            continue

        if _is_company_header(p):
            # Match to Claude's experience
            current_bullets = None
            for co_key, bullets in exp_map.items():
                if co_key in text.lower():
                    current_bullets = bullets
                    break
            skip_title = True
            continue

        if current_bullets is None:
            continue

        # Skip the italic title line (e.g. "Software Engineer Intern, Pixel Recorder Team")
        if skip_title and _is_title_line(p):
            skip_title = False
            continue
        skip_title = False

        # Bullet paragraph: replace with next tailored bullet
        if text and current_bullets:
            new_bullet = current_bullets.pop(0).lstrip("•").strip()
            _replace_paragraph_text(p, new_bullet)

    doc.save(str(output_path))
    return str(output_path)


def save_output(result: dict, job: dict, output_dir: str) -> str:
    """Save tailor suggestions as a .md file (for reference). Returns path."""
    date_str = datetime.now(timezone.utc).strftime("%Y%m%d")
    company = re.sub(r"[^\w\s-]", "", job.get("company", "company")).strip().replace(" ", "_")
    title = re.sub(r"[^\w\s-]", "", job.get("title", "role")).strip().replace(" ", "_")[:40]
    filename = f"{company}_{title}_{date_str}.md"
    output_path = Path(output_dir) / "tailored_resumes" / filename
    output_path.parent.mkdir(parents=True, exist_ok=True)

    lines = [
        f"# Resume Tailoring: {job.get('title')} @ {job.get('company')}",
        f"**Date:** {datetime.now(timezone.utc).strftime('%Y-%m-%d')}",
        f"**URL:** {job.get('url', 'N/A')}",
        "",
        "## Missing Keywords to Add",
        "",
    ]
    for kw in result.get("missing_keywords", []):
        lines.append(f"- `{kw}`")

    lines += ["", "## Papers / Projects to Highlight", ""]
    for paper in result.get("papers_to_highlight", []):
        lines.append(f"- {paper}")

    lines += [
        "",
        "## Suggested Resume Summary",
        "",
        result.get("summary_suggestion", ""),
        "",
        "## Tailored Bullet Points",
        "",
    ]
    for bullet in result.get("tailored_bullets", []):
        lines.append(bullet)

    output_path.write_text("\n".join(lines), encoding="utf-8")
    return str(output_path)
